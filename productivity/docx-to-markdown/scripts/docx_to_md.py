#!/usr/bin/env python3
"""
DOCX → Markdown 转换器
专为中文技术文档设计，保留：标题层级、表格、列表、段落格式（粗体/斜体）

核心策略：
- 标题识别：所有段落样式均为 None，通过正则匹配编号模式识别 H1-H4
- 列表识别：解析 XML 中 <w:numPr> 标签，从 numbering.xml 获取格式，连续同格式项分组编号
- 图片处理：从 document.xml 收集 drawing rId，从 rels 获取实际文件名，按出现顺序编号插入

修复 v2：列表项不再依赖 style.name，直接解析 XML numPr
修复 v3：加入图片插入逻辑（图片跟在所属段落末尾）
"""

import sys
import re
import zipfile
import os
import shutil
from typing import Optional, Tuple, List, Dict
from lxml import etree
from docx import Document


# ---------------------------------------------------------------
# 图片关系映射（从 rels 解析）
# ---------------------------------------------------------------

def load_image_rel_map(docx_path: str) -> Dict[str, str]:
    """从 word/_rels/document.xml.rels 加载 rId -> media文件名 的映射。"""
    with zipfile.ZipFile(docx_path) as z:
        with z.open('word/_rels/document.xml.rels') as f:
            rels_content = f.read().decode('utf-8')

    rel_map: Dict[str, str] = {}
    for match in re.finditer(r'Id="(rId\d+)"[^>]+Target="media/([^"]+)"', rels_content):
        rel_map[match.group(1)] = match.group(2)
    return rel_map


# ---------------------------------------------------------------
# 编号格式映射（从 numbering.xml 解析）
# ---------------------------------------------------------------

def load_numbering_map(docx_path: str) -> Tuple[Dict[str, Tuple[str, str]], Dict[str, str]]:
    """加载 numbering.xml，返回:
    - numId -> (numFmt, lvlText)  如 "81" -> ("bullet", "•")
    """
    with zipfile.ZipFile(docx_path) as z:
        with z.open('word/numbering.xml') as f:
            numbering = f.read().decode('utf-8')

    abstract_map: Dict[str, Tuple[str, str]] = {}
    for abs_id, abs_content in re.findall(
            r'<w:abstractNum w:abstractNumId="(\d+)">(.*?)</w:abstractNum>',
            numbering, re.DOTALL):
        lvl_match = re.search(r'<w:lvl>(.*?)</w:lvl>', abs_content, re.DOTALL)
        if lvl_match:
            lvl_body = lvl_match.group(1)
            lvl_text_match = re.search(r'<w:lvlText[^>]+w:val="([^"]*)"', lvl_body)
            num_fmt_match = re.search(r'<w:numFmt[^>]+w:val="([^"]*)"', lvl_body)
            if lvl_text_match and num_fmt_match:
                abstract_map[abs_id] = (num_fmt_match.group(1), lvl_text_match.group(1))

    num_to_abstract: Dict[str, str] = {}
    for num_id, abs_id in re.findall(
            r'<w:num w:numId="(\d+)"><w:abstractNumId w:val="(\d+)"/></w:num>',
            numbering):
        num_to_abstract[num_id] = abs_id

    num_id_map: Dict[str, Tuple[str, str]] = {}
    for num_id, abs_id in num_to_abstract.items():
        info = abstract_map.get(abs_id)
        if info:
            num_id_map[num_id] = info

    return num_id_map, num_to_abstract


def has_numpr(para) -> Optional[str]:
    """检查段落是否有 numPr，返回 numId 或 None。"""
    numpr = para._element.find(
        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
    if numpr is not None:
        numId_elem = numpr.find(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
        if numId_elem is not None:
            return numId_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
    return None


# ---------------------------------------------------------------
# 列表前缀生成
# ---------------------------------------------------------------

def make_list_prefix(fmt: str, index: int) -> str:
    """根据格式和序号生成 Markdown 列表前缀。"""
    if fmt == 'bullet':
        return '•'
    elif fmt == 'decimal':
        return f'{index}.'
    elif fmt == 'lowerLetter':
        return f'{chr(ord("a") + index - 1)}.'
    elif fmt == 'lowerRoman':
        roman_map = [(10, 'x'), (9, 'ix'), (5, 'v'), (4, 'iv'), (1, 'i')]
        n = index
        result = ''
        for val, sym in roman_map:
            while n >= val:
                result += sym
                n -= val
        return result + '.'
    else:
        return f'{index}.'


def group_and_prefix_list(items: List[Tuple[int, str, str, Optional[str]]]) -> List[Tuple[int, str]]:
    """将连续同格式且同章节上下文的列表项分组并编号。

    Args:
        items: [(elem_idx, text, fmt, section_ctx)] 按文档顺序排列的列表项
               section_ctx 是当前列表项所属的最深章节号，如 "2.8.1" 或 "3.1.3"
    Returns:
        [(elem_idx, prefixed_text)]
    """
    if not items:
        return []

    result = []
    i = 0
    while i < len(items):
        start = i
        fmt = items[i][2]
        ctx = items[i][3] if len(items[i]) > 3 else None
        while i < len(items) and items[i][2] == fmt:
            cur_ctx = items[i][3] if len(items[i]) > 3 else None
            if cur_ctx != ctx:
                break
            i += 1
        end = i
        for j in range(start, end):
            idx = j - start + 1
            prefix = make_list_prefix(fmt, idx)
            result.append((items[j][0], f'{prefix} {items[j][1]}'))
    return result


# ---------------------------------------------------------------
# 标题层级识别
# ---------------------------------------------------------------
HEADING_PATTERN = re.compile(r'^(\d+(?:\.\d+)*)\.?\s+(.+)$')

def get_heading_level(text: str, forbidden_numbers: set = None) -> Optional[Tuple[int, str]]:
    """识别章节标题及层级。返回 (level, title) 或 None。

    Args:
        text: 段落文本
        forbidden_numbers: 当前章节已作为列表编号使用过的数字集合，
                          如 {"1","2","3"} —— 这些数字不应被当作章节标题。
                          例如 "6. ZCUR..." 中 6 若在列表中已出现，则不认作标题。
    """
    m = HEADING_PATTERN.match(text.strip())
    if not m:
        return None
    number_part = m.group(1)
    title_text = m.group(2).strip()
    skip_keywords = [
        "版本号", "修订人", "修订日期", "修订内容",
        "ECU最大响应时间", "Function ID", "Action ID",
        "Bit Offset", "Bit Value", "描述", "建议处理方式", "提示语参考",
    ]
    for kw in skip_keywords:
        if title_text.startswith(kw):
            return None
    parts = [p for p in number_part.split(".") if p]
    if not parts:
        return None

    # 单级编号在子章节上下文中检查 forbidden
    if len(parts) == 1 and forbidden_numbers is not None and number_part in forbidden_numbers:
        return None

    # 启发式：单级编号 + 标题文本较长 → 倾向于是列表项而非章节标题
    # 章节标题通常简短（10字符内），列表项往往是完整描述句
    if len(parts) == 1 and len(title_text) > 18:
        return None

    level = {1: 1, 2: 2, 3: 3}.get(len(parts), 4 if len(parts) >= 4 else 1)
    return level, title_text


# ---------------------------------------------------------------
# 文本格式化
# ---------------------------------------------------------------

def get_paragraph_text(para, strip_formatting: bool = False) -> str:
    """获取段落文本。"""
    result = ""
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if not strip_formatting:
            if run.font.underline:
                text = f"<u>{text}</u>"
            if run.font.strike:
                text = f"~~{text}~~"
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
        result += text
    return result


def para_has_drawing(para) -> bool:
    """检查段落是否包含图片（drawing）。"""
    xml = etree.tostring(para._element, encoding='unicode')
    return '<w:drawing>' in xml


# ---------------------------------------------------------------
# 表格转换
# ---------------------------------------------------------------

def extract_table_markdown(table) -> str:
    """将 Word 表格转为 Markdown 表格"""
    rows = table.rows
    if not rows:
        return ""
    col_count = max(len(row.cells) for row in rows)
    lines = []
    for row_idx, row in enumerate(rows):
        cells = row.cells
        cell_texts = []
        for cell in cells:
            cell_text = cell.text.replace("\r", "").replace("|", "\\|").replace("\n", "<br>")
            cell_texts.append(cell_text)
        while len(cell_texts) < col_count:
            cell_texts.append("")
        cell_texts = cell_texts[:col_count]
        if row_idx == 0:
            lines.append("| " + " | ".join(cell_texts) + " |")
            lines.append("| " + " | ".join(["---"] * col_count) + " |")
        else:
            lines.append("| " + " | ".join(cell_texts) + " |")
    return "\n".join(lines) + "\n"


def build_table_map(doc: Document):
    """建立 table element id → table 对象的映射"""
    return {id(t._element): t for t in doc.tables}


# ---------------------------------------------------------------
# 主转换逻辑
# ---------------------------------------------------------------

def convert_docx_to_markdown(doc_path: str, output_path: str, images_dir: str = ""):
    """
    将 DOCX 转换为 Markdown。

    Args:
        doc_path: 输入 .docx 路径
        output_path: 输出 .md 路径
        images_dir: 图片提取目录（可选）
    """
    doc = Document(doc_path)
    table_map = build_table_map(doc)
    elem_list = list(doc.element.body)

    # 加载编号映射
    num_id_map, _ = load_numbering_map(doc_path)

    # 收集图片信息
    rel_map = load_image_rel_map(doc_path)

    # 收集所有带 drawing 的段落 elem_idx，并分配图片编号
    # 图片跟在 drawing 所在段落之后、空行之后输出
    drawing_info: List[Tuple[int, str, int]] = []  # [(elem_idx, rel_id, image_num)]
    image_num = 0
    para_pattern = re.compile(r'<w:p\b[^>]*>(.*?)</w:p>', re.DOTALL)

    # 建立 elem -> idx 映射
    elem_to_idx = {id(elem): idx for idx, elem in enumerate(elem_list)}

    for para in doc.paragraphs:
        xml = etree.tostring(para._element, encoding='unicode')
        if '<w:drawing>' not in xml:
            continue
        blip_match = re.search(r'r:embed="(rId\d+)"', xml)
        if not blip_match:
            continue
        rId = blip_match.group(1)
        if rId not in rel_map:
            continue
        elem_idx = elem_to_idx.get(id(para._element))
        if elem_idx is not None:
            image_num += 1
            drawing_info.append((elem_idx, rId, image_num))

    # 追踪当前章节上下文：用于列表项分组
    # 第一次扫描：建立 section_at_idx（elem_idx -> section）
    # heading 段落使用更新前的 current_section（父章节）
    section_at_idx: Dict[int, str] = {}
    current_section = ""
    HEADING_NUM_PATTERN = re.compile(r'^(\d+(?:\.\d+)*)\.?\s+(.+)$')
    for elem_idx, elem in enumerate(elem_list):
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            continue
        para = None
        for p in doc.paragraphs:
            if p._element is elem:
                para = p
                break
        if para is None:
            continue
        text = get_paragraph_text(para).strip()
        m = HEADING_NUM_PATTERN.match(text)
        if m:
            section_at_idx[elem_idx] = current_section
            current_section = m.group(1)
        else:
            section_at_idx[elem_idx] = current_section

    # 收集所有列表项 (elem_idx, text, fmt, section_ctx)
    all_list_items: List[Tuple[int, str, str, str]] = []
    for elem_idx, elem in enumerate(elem_list):
        tag = elem.tag.split('}')[-1]
        if tag != "p":
            continue
        para = None
        for p in doc.paragraphs:
            if p._element is elem:
                para = p
                break
        if para is None:
            continue
        text = get_paragraph_text(para).strip()
        if not text:
            continue
        num_id = has_numpr(para)
        if num_id:
            info = num_id_map.get(num_id, ('unknown', '?'))
            fmt = info[0]
            ctx = section_at_idx.get(elem_idx, "")
            all_list_items.append((elem_idx, text, fmt, ctx))

    # 列表分组编号
    list_fixes = group_and_prefix_list(all_list_items)
    list_output_map: Dict[int, str] = {item[0]: item[1] for item in list_fixes}

    # 预先建立每个章节的 forbidden numbers 集合（仅处理纯文本编号）
    # 规则：只禁止在子章节（H2+）上下文中出现的 "N. xxx" 文本，
    #       因为这些可能是有序列表而非子章节标题（如 PE 流程步骤 1-6）
    forbidden_per_section: Dict[str, set] = {}
    LIST_NUM_PATTERN = re.compile(r'^(\d+)\.')

    # 扫描所有段落，收集以 "N. " 开头但没有 numPr 的纯文本编号
    for elem_idx, elem in enumerate(elem_list):
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            continue
        ctx = section_at_idx.get(elem_idx, "")
        # 只处理子章节（H2+）上下文；顶层和 H1 章节不需要禁止
        if not ctx or ctx.count('.') == 0:
            continue
        para = None
        for p in doc.paragraphs:
            if p._element is elem:
                para = p
                break
        if para is None:
            continue
        text = get_paragraph_text(para).strip()
        if not text:
            continue
        if has_numpr(para):  # 有 numPr 的列表项不参与 forbidden
            continue
        m = LIST_NUM_PATTERN.match(text)
        if m:
            num = m.group(1)
            if ctx not in forbidden_per_section:
                forbidden_per_section[ctx] = set()
            forbidden_per_section[ctx].add(num)

    # 建立 drawing 信息：elem_idx -> image_markdown_line
    drawing_output_map: Dict[int, str] = {}
    for elem_idx, rId, image_num in drawing_info:
        filename = rel_map[rId]
        ext = os.path.splitext(filename)[1] or '.png'
        img_basename = os.path.basename(images_dir) if images_dir else 'images'
        drawing_output_map[elem_idx] = f"![image{image_num}{ext}]({images_dir}/image{image_num}{ext})"

    # 第二次扫描：生成 md 内容
    md_lines: List[str] = []
    blank_count = 0
    prev_was_heading = False
    current_section = ""  # 当前章节号

    for elem_idx, elem in enumerate(elem_list):
        tag = elem.tag.split('}')[-1]

        # 图片段落：drawing 在空文本段落中，优先处理
        if tag == 'p' and elem_idx in drawing_output_map:
            md_lines.append('\n')
            md_lines.append(drawing_output_map[elem_idx] + '\n')
            md_lines.append('\n')
            blank_count = 0
            prev_was_heading = False
            continue

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is elem:
                    para = p
                    break
            if para is None:
                continue

            text = get_paragraph_text(para).strip()

            # 列表项
            if elem_idx in list_output_map:
                md_lines.append(list_output_map[elem_idx] + "\n")
                blank_count = 0
                prev_was_heading = False
                if elem_idx in drawing_output_map:
                    md_lines.append(drawing_output_map[elem_idx] + "\n")
                    md_lines.append("\n")
                continue

            # 空段落
            if not text:
                blank_count = 0
                continue

            # 文档第一行（标题）
            if para is doc.paragraphs[0]:
                md_lines.append(text + "\n")
                blank_count = 0
                prev_was_heading = False
                continue

            _dbg_fname = forbidden_per_section.get(section_at_idx.get(elem_idx, ""), set())
            heading_info = get_heading_level(
                text,
                _dbg_fname
            )
            if "2.6" in text or "2.5" in text or "2.7" in text:
                import sys as _sys
                _sys.stderr.write(f'DBG elem_idx={elem_idx} section={section_at_idx.get(elem_idx,"?")} forbidden={_dbg_fname} text={repr(text[:80])} heading_info={heading_info}\n')
                _sys.stderr.flush()
            if heading_info:
                level, title = heading_info
                md_lines.append("#" * level + " " + title + "\n")
                blank_count = 0
                prev_was_heading = True
                if "2.6" in text or "2.5" in text or "2.7" in text:
                    import sys as _sys
                    _sys.stderr.write(f'DBG HEADING: elem_idx={elem_idx} heading_info={heading_info}\n')
                    _sys.stderr.flush()
            else:
                md_lines.append(text + "\n")
                blank_count = 0
                prev_was_heading = False
                if "2.6" in text or "2.5" in text or "2.7" in text:
                    import sys as _sys
                    _sys.stderr.write(f'DBG NON-HEADING: elem_idx={elem_idx} text={repr(text[:80])}\n')
                    _sys.stderr.flush()

        elif tag == "tbl":
            table = table_map.get(id(elem))
            if table:
                if len(table.rows) == 1 and len(table.rows[0].cells) == 1:
                    for p in table.rows[0].cells[0].paragraphs:
                        t = get_paragraph_text(p).strip()
                        if t:
                            md_lines.append(t + "\n")
                    blank_count = 0
                    prev_was_heading = False
                    continue
                md_lines.append(extract_table_markdown(table))
                md_lines.append("\n")
                blank_count = 0
                prev_was_heading = False

        elif tag == "bookmarkStart":
            continue

    # 清理连续空行（最多保留2个）
    cleaned = []
    blank_count = 0
    for line in md_lines:
        if not line.strip():
            blank_count += 1
            if blank_count <= 2:
                cleaned.append("")
        else:
            blank_count = 0
            cleaned.append(line.rstrip())

    md_content = "\n".join(cleaned)

    # 去掉标题中多余的 ** 包裹
    lines = md_content.split("\n")
    fixed = []
    for line in lines:
        m = re.match(r'^(#{1,6})\s+\*\*(.+)\*\*$', line)
        if m:
            line = m.group(1) + " " + m.group(2)
        m2 = re.match(r'^\*\*T68.+?\*\*$', line)
        if m2:
            line = line.strip("*")
        fixed.append(line)
    md_content = "\n".join(fixed)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    # 提取图片
    if images_dir:
        extract_images(doc_path, images_dir, rel_map)

    # 统计
    h1 = sum(1 for l in md_lines if l.startswith("# ") and not l.startswith("## "))
    h2 = sum(1 for l in md_lines if l.startswith("## ") and not l.startswith("### "))
    h3 = sum(1 for l in md_lines if l.startswith("### ") and not l.startswith("#### "))
    h4 = sum(1 for l in md_lines if l.startswith("#### "))
    table_blocks = sum(1 for l in md_lines if l.startswith("|") and "---" not in l)
    list_items = sum(1 for l in md_lines if l.strip().startswith(("•", "1.", "a.", "i.")))
    image_refs = sum(1 for l in md_lines if l.strip().startswith("![image"))

    print(f"转换完成：{output_path}")
    print(f"H1: {h1}, H2: {h2}, H3: {h3}, H4: {h4}, 表格: {table_blocks}, 列表项: {list_items}, 图片: {image_refs}")
    print(f"总字符: {len(md_content):,}, 总行数: {md_content.count(chr(10)) + 1:,}")


def extract_images(docx_path: str, output_dir: str, rel_map: Dict[str, str] = None):
    """从 docx 中提取所有图片到指定目录，按出现顺序命名 imageN.ext。"""
    os.makedirs(output_dir, exist_ok=True)

    if rel_map is None:
        rel_map = load_image_rel_map(docx_path)

    # rel_map: rId -> media filename (e.g., "image1.png")
    # 按 rId 数字排序后按顺序编号
    rid_to_file = {}
    for rid, filename in rel_map.items():
        # 提取 rId 数字
        m = re.search(r'rId(\d+)', rid)
        if m:
            rid_num = int(m.group(1))
            rid_to_file[rid_num] = filename

    with zipfile.ZipFile(docx_path) as z:
        image_files = [f for f in z.namelist() if f.startswith('word/media/')]
        # 按 rId 数字顺序分配编号
        for idx, (rid_num) in enumerate(sorted(rid_to_file.keys()), 1):
            filename = rid_to_file[rid_num]
            src_path = f'word/media/{filename}'
            ext = os.path.splitext(filename)[1] or '.png'
            out_name = f"image{idx}{ext}"
            out_path = os.path.join(output_dir, out_name)
            try:
                with z.open(src_path) as src, open(out_path, 'wb') as dst:
                    shutil.copyfileobj(src, dst)
            except KeyError:
                pass

    print(f"图片已提取: {len(rid_to_file)} 张 -> {output_dir}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python docx_to_md.py <输入.docx> <输出.md> [图片目录]")
        sys.exit(1)
    img_dir = sys.argv[3] if len(sys.argv) > 3 else ""
    convert_docx_to_markdown(sys.argv[1], sys.argv[2], img_dir)
